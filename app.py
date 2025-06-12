from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openai
import os
import tempfile

# ── Flask setup ────────────────────────────────────────────────────────────────
app = Flask(__name__)

# ── OpenAI setup ───────────────────────────────────────────────────────────────
openai.api_key = os.getenv("OPENAI_API_KEY")   # <-- Set this in Render dashboard

# ── GPT: return structured Markdown (no visual styling) ───────────────────────
def gpt_format_text(raw_text: str) -> str:
    system_prompt = (
        "You are a professional technical writer. "
        "Rewrite input text with clear Markdown headers (#, ##, ###), bullet lists (- item) "
        "and short paragraphs."
    )
    user_prompt = f"Format the following into structured Markdown:\n\n\"\"\"\n{raw_text}\n\"\"\""
    response = openai.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": system_prompt},
                  {"role": "user", "content": user_prompt}],
        temperature=0.5,
    )
    return response.choices[0].message.content.strip()

# ── DOCX helpers ───────────────────────────────────────────────────────────────
def add_page_numbers(section):
    """Center-aligned footer with automatic PAGE field."""
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()

    # Field code: PAGE
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText');   instr.text = ' PAGE '
    fld_end  = OxmlElement('w:fldChar');  fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

def generate_docx(markdown_text: str, filename: str) -> str:
    doc = Document()

    # Page margins
    sect = doc.sections[0]
    sect.top_margin    = sect.bottom_margin = Inches(1)
    sect.left_margin   = sect.right_margin  = Inches(1)
    add_page_numbers(sect)

    # Title (first line with single '#')
    lines = markdown_text.splitlines()
    if lines and lines[0].startswith("# "):
        title_text = lines.pop(0)[2:].strip()
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(22)
        title.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)  # blue
        doc.add_paragraph("")  # spacer

    for line in lines:
        stripped = line.strip()

        # Headings
        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=1)
            h.runs[0].font.color.rgb = RGBColor(0x00, 0x57, 0xA6)  # darker blue
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=2)
            h.runs[0].font.color.rgb = RGBColor(0x00, 0x57, 0xA6)
        # Bullets
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        # Blank line ⇒ extra spacing
        elif stripped == "":
            doc.add_paragraph("")
        # Normal paragraph
        else:
            p = doc.add_paragraph(stripped)
            p.style.font.size = Pt(11)

    doc.save(filename)
    return filename

# ── Route ──────────────────────────────────────────────────────────────────────
@app.route("/generate-docx", methods=["POST"])
def generate_docx_endpoint():
    data = request.get_json(silent=True)
    if not data or "text" not in data:
        return jsonify({"error": "Provide JSON with a 'text' field"}), 400

    try:
        markdown = gpt_format_text(data["text"])
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            generate_docx(markdown, tmp.name)
            return send_file(
                tmp.name,
                as_attachment=True,
                download_name="ai_document.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

@app.route("/")
def health():
    return "✅ DOCX API is up"

if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
